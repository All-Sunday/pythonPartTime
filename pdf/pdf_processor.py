# @Description: TODO
# @Author https://github.com/All-Sunday All-Sunday
# @Time 2022/4/10 15:38
# @File : pdf_processor.py
import os

import docx
from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LTTextBoxHorizontal, LAParams
from nltk.tokenize import sent_tokenize


def parse(DataIO):
    # 接受解析结果字符串
    result = ''
    # 通过传入的pdf文件对象 创建pdf文档分析器
    pdf_parser = PDFParser(DataIO)
    # 创建pdf文档
    pdf_doc = PDFDocument()
    # 分析器、文档连接
    pdf_parser.set_document(pdf_doc)
    pdf_doc.set_parser(pdf_parser)
    # 初始化文档
    pdf_doc.initialize()

    # 创建管理共享资源的pdf资源管理器
    pdf_resource = PDFResourceManager()
    # 创建参数分析器
    pdf_laparams = LAParams()
    # 创建聚合器
    device = PDFPageAggregator(pdf_resource, laparams=pdf_laparams)
    # 创建页面解释器
    interpreter = PDFPageInterpreter(pdf_resource, device)

    for page in pdf_doc.get_pages():
        # 页面解析器来读取
        interpreter.process_page(page)
        # 聚合器来获取内容
        layout = device.get_result()
        # layout是一个LTPage对象 存放page解析出来的各种对象 有LTTextBox，LTImage，LTTextBoxHorizontal 等

        for x in layout:
            try:
                if isinstance(x, LTTextBoxHorizontal):
                    # 获取文本 拼接结果
                    result += x.get_text()
            except:
                print("Failed")
    return result


def processor(path):
    filelist = []
    res = ''
    # 获取path路径下的所有pdf文件  包括子文件夹
    for root, dirs, files in os.walk(path):
        for name in files:
            if name[-3:] == 'pdf':
                pdf_path = path + name
                filelist.append(pdf_path)
                # 解析pdf文件对象
                print(pdf_path)
                with open(pdf_path, 'rb') as pdf:
                    res += parse(pdf)
    print(filelist)
    return res


if __name__ == '__main__':

    # path = r'\'
    path = input('请输入文件夹路径：')

    # 获取输入单词的list
    target_list = input('请输入要提取的单词，以空格分隔：').split(' ')
    # target_list = ['art', 'religion', 'happiness', 'argue', 'figure']

    # 将list转为dict  每个单词对应一个空list 用于接受结果
    # target = {'art': [], 'religion': [], 'happiness': [], 'argue':[], 'figure':[]}
    target_dict = {}
    for i in target_list:
        target_dict[i] = []
    print(target_dict)

    # 解析path路径下的所有pdf 结果是str
    res = processor(path)
    # 通过ntlk进行英语分句 结果是list
    res = sent_tokenize(res)

    # year记录结果的年份
    year = 2004
    # 遍历结果list
    for i in res:
        # 如果遇到 年全国硕士研究生招生考试 说明是新的一年的卷子
        if '年全国硕士研究生招生考试' in i:
            year += 1

        # 遍历目标单词dict  如果单词出现在句子里 则加入单词对应的列表
        for k, l in target_dict.items():
            # if (' art ' in i) or (' art.' in i):
            if (' ' + k + ' ' in i) or (' ' + k + '.' in i):
                # 如果符合目标单词格式 则加入到单词对应的list 并追加年份
                l.append(i + '（' + str(year) + '）')
    # print(target_dict, len(target_dict))

    # 创建docx对象
    docx_file = docx.Document()
    # 遍历单词dict  写入结果到docx
    for k, l in target_dict.items():
        # 先写入单词 及相应的个数
        docx_file.add_paragraph(k + ' ' + len(l).__str__())
        # num记录结果序号
        num = 1
        for i in l:
            # 句子中会出现换行  进行替换
            docx_file.add_paragraph(str(num) + '. ' + i.replace('\n', ''))
            num += 1
        docx_file.add_paragraph('__________________________________________________________________________________')
    # 保存
    docx_file.save("result.docx")
