# @Description: TODO
# @Author https://github.com/All-Sunday All-Sunday
# @Time 2022/4/10 15:38
# @File : pdf_processor.py
import re
import time
import urllib
import importlib, sys

import docx

importlib.reload(sys)
from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfdevice import PDFDevice
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LTTextBoxHorizontal, LAParams
from pdfminer.pdfinterp import PDFTextExtractionNotAllowed
from nltk.tokenize import sent_tokenize


def parse(DataIO):
    result = ''
    # 用文件对象创建一个PDF文档分析器
    parser = PDFParser(DataIO)
    # 创建一个PDF文档
    doc = PDFDocument()
    # 分析器和文档相互连接
    parser.set_document(doc)
    doc.set_parser(parser)
    # 提供初始化密码，没有默认为空
    doc.initialize()
    # 检查文档是否可以转成TXT，如果不可以就忽略

    # 创建PDF资源管理器，来管理共享资源
    rsrcmagr = PDFResourceManager()
    # 创建一个PDF设备对象
    laparams = LAParams()
    # 将资源管理器和设备对象聚合
    device = PDFPageAggregator(rsrcmagr, laparams=laparams)
    # 创建一个PDF解释器对象
    interpreter = PDFPageInterpreter(rsrcmagr, device)

    # 循环遍历列表，每次处理一个page内容
    # doc.get_pages()获取page列表

    for page in doc.get_pages():
        interpreter.process_page(page)
        # 接收该页面的LTPage对象
        layout = device.get_result()
        # 这里的layout是一个LTPage对象 里面存放着page解析出来的各种对象
        # 一般包括LTTextBox，LTFigure，LTImage，LTTextBoxHorizontal等等一些对像
        # 想要获取文本就得获取对象的text属性
        for x in layout:
            try:
                if (isinstance(x, LTTextBoxHorizontal)):
                    result += x.get_text()
                    # print(result)

            except:
                print("Failed")
    return result


if __name__ == '__main__':
    res = ''
    # 解析本地PDF文本，保存到本地TXT
    with open(r'2005-2009真题(英语一二通用).pdf', 'rb') as pdf:
        res += parse(pdf)
    with open(r'2010-2021英语一真题.pdf', 'rb') as pdf:
        res += parse(pdf)

    # list_ret = list()
    #
    # for s_str in res.split('.'):   #对输入进行处理  (用英文结尾句号.来划分句子)
    #     s_str = s_str.replace('\n','')      #去掉句子中的\n换行
    #
    #     if '?' in s_str:
    #         list_ret.extend(s_str.split('?'))
    #     elif '!' in s_str:
    #         list_ret.extend(s_str.split('!'))
    #     else:
    #         list_ret.append(s_str)
    #
    # for s_str in list_ret:
    #     #print(s_str+".\n")
    #     s_str=s_str+".\n"         #每一个完整英语句子加上句号“.”，然后加个换行
    #     print(s_str)
    # res = "A pig."
    # pattern = re.compile(r'[A-Za-z]+[A-Za-z0-9_,\"#;.() \\s]*[.]$')   # 查找数字
    # result1 = pattern.findall(res, re.S)
    # print(result1)
    # result1 = pattern.search(res, re.S)
    # print(result1)
    # result1 = pattern.match(res, re.S)
    # print(result1)

    res = sent_tokenize(res)
    # print(res)
    target_list = input('请输入要提取的单词，以空格分隔：').split(' ')
    # print(target_list)

    # target_list = ['art', 'religion', 'happiness', 'argue', 'figure']
    target_dict = {}
    for i in target_list:
        target_dict[i] = []
    print(target_dict)
    # time.sleep(100)
    # target = {'art': [], 'religion': [], 'happiness': [], 'argue':[], 'figure':[]}
    year = 2004
    for i in res:
        if '年全国硕士研究生招生考试' in i:
            year += 1
        # if (' art ' in i) or (' art.' in i):
        for k, l in target_dict.items():
            if (' ' + k + ' ' in i) or (' ' + k + '.' in i):
                l.append(i + '（' + str(year) + '）')
    print(target_dict, len(target_dict))

    file = docx.Document()

    # for i in target:
    #     file.add_paragraph(i.replace('\n', ''))
    for k, l in target_dict.items():
        file.add_paragraph(k + ' ' + len(l).__str__())
        num = 1
        for i in l:
            file.add_paragraph(str(num) + '. '  + i.replace('\n', ''))
            # file.add_paragraph(i.replace('\n', ''))
            num += 1
        file.add_paragraph('__________________________________________________________________________________')
    file.save("result.docx")
